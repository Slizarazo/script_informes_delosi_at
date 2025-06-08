import docx
import matplotlib.pyplot as plt
import io
from docx.shared import Inches

def replace_text_with_chart(doc_path, replacements):
    """Reemplaza fragmentos de texto en un documento Word por gráficos generados."""
    doc = docx.Document(doc_path)
    
    for para in doc.paragraphs:
        for text_to_replace, chart_func in replacements.items():
            if text_to_replace in para.text:
                para.text = para.text.replace(text_to_replace, "")
                image_stream = chart_func()
                run = para.add_run()
                run.add_picture(image_stream, width=Inches(3))
    
    output_path = "Prueba_informe.docx"
    doc.save(output_path)
    print(f"Documento guardado como {output_path}")
    return output_path

def generate_chart():
    """Genera una gráfica simple y la devuelve como un objeto BytesIO."""
    fig, ax = plt.subplots()
    ax.plot([1, 2, 3, 4], [10, 20, 25, 30], marker='o')
    ax.set_title("Ejemplo de Gráfica")
    
    image_stream = io.BytesIO()
    plt.savefig(image_stream, format='png')
    image_stream.seek(0)
    return image_stream

# Uso
replacements = {
    "{{GRAFICA_1}}": generate_chart
}

# Llamar a la función con un archivo .docx existente
replace_text_with_chart("test_informe.docx", replacements)
