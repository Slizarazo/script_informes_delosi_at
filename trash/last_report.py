import re
from docx import Document
from docx2pdf import convert
import os

def reemplazar_llaves_y_exportar_pdf(path_docx, reemplazos, output_pdf_path=None):
    # Abrir el documento
    doc = Document(path_docx)

    # Función auxiliar para reemplazar texto dentro de un párrafo
    def reemplazar_en_parrafo(parrafo, reemplazos):
        for key, val in reemplazos.items():
            llave = f"[[{key}]]"
            if llave in parrafo.text:
                # Reemplazo manteniendo formato original
                for run in parrafo.runs:
                    if llave in run.text:
                        run.text = run.text.replace(llave, val)

    # Recorremos párrafos y tablas
    for parrafo in doc.paragraphs:
        reemplazar_en_parrafo(parrafo, reemplazos)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_en_parrafo(parrafo, reemplazos)

    # Guardar documento temporal modificado
    temp_docx = "documento_modificado.docx"
    doc.save(temp_docx)

    # Convertir a PDF
    if output_pdf_path is None:
        output_pdf_path = os.path.splitext(path_docx)[0] + "_modificado.pdf"
    convert(temp_docx, output_pdf_path)

    # Opcional: eliminar el archivo docx temporal
    os.remove(temp_docx)

    return output_pdf_path


reemplazos = {
    "Nombre": "Juan Perez",
    "Fecha_assesment": "27 de abril del 2025"
}

ruta_pdf = reemplazar_llaves_y_exportar_pdf("inf_delosi.docx", reemplazos)
print(f"PDF generado: {ruta_pdf}")
