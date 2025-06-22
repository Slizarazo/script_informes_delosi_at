import re
from docx import Document
from docx2pdf import convert
import os, sys
import matplotlib.pyplot as plt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from graficos import *

# ╔════════════════════════════════════════════════════════════════════╗
# ║                            ⚠️  ATENCIÓN                            ║
# ╠════════════════════════════════════════════════════════════════════╣
# ║                                                                    ║
# ║   ESTE ARCHIVO ES CRÍTICO PARA EL FUNCIONAMIENTO DEL SISTEMA.     ║
# ║   NO MODIFICAR NI ELIMINAR SU CONTENIDO SIN AUTORIZACIÓN.         ║
# ║                                                                    ║
# ║   Cualquier cambio podría generar errores graves en la aplicación.║
# ║                                                                    ║
# ╚════════════════════════════════════════════════════════════════════╝

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), 'config')))

from config.models import Dim_evaluated

def reemplazar_llaves_y_exportar_pdf(path_docx, reemplazos, output_doc, marca, graficos={}, output_pdf_path=None):
    doc = Document(path_docx)
    
    def reemplazar_o_insertar(parrafo, reemplazos, graficos):
        for key, val in reemplazos.items():
            llave = f"[[{key}]]"
            if llave in parrafo.text:
                for run in parrafo.runs:
                    if llave in run.text:
                        run.text = run.text.replace(llave, val)

        for key, ruta_img in graficos.items():
            llave = f"[[{key}]]"
            if llave in parrafo.text:
                # Eliminar texto marcador y reemplazarlo con imagen
                for run in parrafo.runs:
                    if llave in run.text:
                        run.text = run.text.replace(llave, '')
                run = parrafo.add_run()
                if key == "Grafico_1":
                    ancho, alto = Inches(8), Inches(2.5)
                elif key == "Grafico_2":
                    ancho, alto = Inches(8), Inches(2.8)
                elif key == "Tabla_1":
                     ancho, alto = Inches(8), Inches(4.5)
                else:
                    ancho, alto = Inches(8), Inches(1.5)

                run.add_picture(ruta_img, ancho, alto)

                parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for parrafo in doc.paragraphs:
        reemplazar_o_insertar(parrafo, reemplazos, graficos)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_o_insertar(parrafo, reemplazos, graficos)

    temp_docx = "documento_modificado.docx"
    doc.save(temp_docx)

    # Crear carpeta si no existe
    ruta_base = "pdfs"  # Sin "/" inicial para que sea relativa al proyecto
    ruta_completa = os.path.join(ruta_base, marca)

    os.makedirs(ruta_completa, exist_ok=True)

    if output_pdf_path is None:
        output_pdf_path = output_doc + ".pdf"

    output_pdf_path = os.path.join('pdfs/'+marca, output_pdf_path)
    convert(temp_docx, output_pdf_path)
    os.remove(temp_docx)

    return output_pdf_path

base_dir = os.path.dirname(os.path.abspath(__file__))
file_path = os.path.join(base_dir, 'excel', 'delosi_asistentes_tienda_dimensions.xlsx')

df = pd.read_excel(file_path)
data = Dim_evaluated.get_all()

for reg in data:

    marca = str(reg['marca'].replace(' ', '_'))
    usuario = str(reg['nombre_display'])

    # Ruta base
    base_path = "graficas/calificacion_global"
    # Nombre de la subcarpeta según el registro
    nombre_carpeta = marca if marca else "Sin Marca"
    # Ruta completa
    carpeta_destino = os.path.join(base_path, nombre_carpeta)

    # Validar existencia y crear si no existe
    if not os.path.exists(carpeta_destino):
        os.makedirs(carpeta_destino)
        print(f"📁 Carpeta creada: {carpeta_destino}")
    else:
        print(f"📂 Carpeta ya existe: {carpeta_destino}")
        
    # Generar imagen de gráfico
    calificacion_global = "graficas/calificacion_global/" + marca + "/" + usuario + ".png"
    drpc = "graficas/distribucion_de_resultados_ponderados_por_capacidad/" + marca + "/" + usuario + ".png"
    drpf = "graficas/distribucion_de_resultados_ponderados_por_fase/" + marca + "/" + usuario + ".png"
    dpcyf = "graficas/distribucion_por_capacidad_y_fase/" + marca + "/" + usuario + ".png"
    comp_resultados = "graficas/comparativa_de_resultados/" + marca + "/" + usuario + ".png"
    
    plot_por_questionblock(df, usuario, marca)
    mostrar_tarjetas(reg['caso_practico'], reg['conocimiento'], usuario, marca)
    plot_ponderado(df[df.evaluatedemployeedisplayname == usuario], usuario, marca)
    crear_tabla_distribucion_v4(df, usuario, marca)
    calcular_y_graficar_tarjetas(df, usuario, marca)
    
    reemplazos = {
        "Nombre": usuario,
        "Fecha_assesment": "28 de Octubre de 2024 a 27 de Enero de 2025",
        "Unidad_negocio": reg['unidad_negocio'],
        "Marca": reg['marca']
    }

    graficos = {
        "Calificacion_global": calificacion_global,
        "Grafico_1": drpc,
        "Grafico_2": drpf,
        "Tabla_1": dpcyf,
        "Comparativa_resultados": comp_resultados
    }

    ruta_pdf = reemplazar_llaves_y_exportar_pdf("inf_delosi.docx", reemplazos, str(reg['nombre']), str(reg['marca']), graficos)
    print(f"PDF generado: {ruta_pdf}")
